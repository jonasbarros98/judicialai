import openai
from django.http import HttpResponse
from xhtml2pdf import pisa
from django.template.loader import get_template
from io import BytesIO
from senha_gpt import API_KEY
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.shortcuts import get_object_or_404
from documentos.models import DocumentoJuridico
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from bs4 import BeautifulSoup  # Usaremos para processar o HTML
from docx.shared import Pt, Inches
from search_web_gpt import buscar_jurisprudencias_bing,gerar_frase_pesquisa_gpt
from django import template
from django.http import StreamingHttpResponse



openai.api_key = API_KEY

def gerar_conteudo_juridico(tipo_documento, dados_preenchimento):
    try:
        # Extração dos dados do caso para a pesquisa de jurisprudência
        tipo_acao = dados_preenchimento.get('tipo_acao')
        valor_causa = dados_preenchimento.get('valor_causa')
        juizo_competente = dados_preenchimento.get('juizo_competente')
        descricao_fatos = dados_preenchimento.get('descricao_fatos')
        dados_requerente = dados_preenchimento.get('dados_requerente') or '[Nome do Requerente]'
        dados_requerido = dados_preenchimento.get('dados_requerido') or '[Nome do Requerido]'
        justica_gratis = dados_preenchimento.get('justica_gratis') or 'Não Solicite Justiça Gratuita'
        
        # Parte 1: Buscar jurisprudências relevantes
        
        query = gerar_frase_pesquisa_gpt(descricao_fatos)
        jurisprudencias = buscar_jurisprudencias_bing(query)
        
        print(f"Gerando frase de pesquisa para o caso: '{query}'..")    

        # Exibir as jurisprudências encontradas (para fins de teste)
        print(f"Jurisprudências encontradas: {jurisprudencias}")
        # Construímos a mensagem para enviar ao modelo, agora incluindo as jurisprudências reais
        mensagem = (
            f"Você é um advogado altamente especializado em direito civil, com mais de 20 anos de experiência, e precisa redigir uma Petição Inicial para uma ação do tipo: {tipo_acao} \n"
            f"O valor da causa é de R$ {valor_causa}, e o juízo competente é {juizo_competente}. \n"
            f"Os fatos principais do caso são os seguintes: {descricao_fatos}. "
            f"Os dados conhecidos do requerente são: {dados_requerente}.\n"
            f"os dados conhecidos do requerido são: {dados_requerido}.\n"
            f"- Quando não forem fornecidos os dados do cliente, empresa ou envolvidos no caso, não invente dados de endereço ou afins, coloque-os entre colchetes aguardando inclusão. Ex: [nome], [endereço]."
            f"\n\nInstruções adicionais:\n"
            f"- Inclua a(s) jurisprudência(s) relevante(s) da pesquisa e cite-as corretamente no formato de peça jurídica em HTML.\n"
            f"- Inclua fundamentação legal apropriada, mencionando artigos do Código Civil, Código de Processo Civil, ou outras legislações pertinentes.\n"
            f"- Caso seja citada a jurisprudencia, utilize formatação em HTML, destaque os trechos de legislação ou jurisprudência em <i>itálico</i> , entre aspas e com espaçamento diferenciado. Exemplo: <i>Art. 373. O ônus da prova incumbe ao autor...</i>."
            f"- Não inclua o link da pesquisa de jurisprudência no documento final. Ao inves disso você deve formatá-la igual uma citação profissional na peça juridica.\n"
            f"- Mantenha a argumentação objetiva e focada no necessário para a defesa ou acusação, evitando argumentos vagos ou irrelevantes.\n"
            f"- Todo o documento deve ser em HTML. Não insira marcadores do tipo ```html no inicio ou qualquer outra forma de marcação de código. Traga apenas a estrutura HTML pura.\n"
            f"- A estrutura deve começar com as tags <html>, <head> e <body>. NÃO gere a tag <tittle>\n"
            f"- Utilize as tags <b>negrito</b> para destacar informações importantes, como títulos de seções, palavras-chave ou valores significativos.\n"
            f"- Evite espaços desnecessários ao gerar o documento.\n"
            f"- A estrutura deve conter seções como: <h3>DOS FATOS</h3>, <h3>DO DIREITO</h3>, e <h3>DO PEDIDO</h3>, etc., numerando subtítulos em algarismos romanos (I, II, III).\n"
            f"- Os titulos do inicio devem ser em <h2>, durante o texto pode utilizar <h3> ou <h4>. Nunca use h1 neste documento."
            f"- Finalize com espaço para nome do advogado, OAB e data.\n"
            f"- O documento deve ter clareza e fluidez, com parágrafos bem organizados e fundamentação robusta. Use a formatação HTML para manter o documento corretamente estruturado.\n"
            f"- O documento completo deve conter pelo menos 2000 palavras."
            f"- O campo a seguir define se o cliente solicita justiça gratuita ou nao: {justica_gratis}.\n"
            f"- As jurisprudências reais encontradas na pesquisa na web são as seguintes abaixo (mas use apenas as que você considerar pertinentes):\n"
              
        )

        # Parte 2: Incluir as jurisprudências reais na mensagem para o GPT-4
        for idx, jurisprudencia in enumerate(jurisprudencias["webPages"]["value"][:3]):
            mensagem += f"{idx + 1}. Título: {jurisprudencia['name']}\n"
            mensagem += f"   URL: {jurisprudencia['url']}\n"
            mensagem += f"   Trecho: {jurisprudencia['snippet']}\n\n"

        mensagem += (
            "\nUse essas informações para compor a melhor citação possível. "
            "Formate a citação em HTML e siga as instruções."
        )


        # Chamamos a API OpenAI usando streaming
        
        response = openai.chat.completions.create(
            model="chatgpt-4o-latest",  # Verifique se o modelo está correto
            messages=[
                {"role": "system", "content": "Você é um advogado especializado em direito civil com 20 anos de experiencia."},
                {"role": "user", "content": mensagem}],
            temperature=0.5,
            max_tokens=3640,
            #stream=False
            stream = True
        )
        
        conteudo_completo = ""
        # Processar a resposta por partes
        for parte in response:
            if not parte.choices[0].delta == {}:
                print(str(parte.choices[0].delta.content), end="")
                conteudo_completo += str(parte.choices[0].delta.content)

        return conteudo_completo

    except Exception as e:
        print("Erro:", e)
        return "Erro: " + str(e)


def render_pdf_view(request, documento):
    try:
        # Carrega o template HTML
        template = get_template('documentos/documento_pdf_template.html')

        # Passa os dados do documento para o template
        html = template.render({'documento': documento})

        # Configura a resposta HTTP para um arquivo PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="documento_{documento.id}.pdf"'

        # Gera o PDF a partir do conteúdo HTML
        pisa_status = pisa.CreatePDF(BytesIO(html.encode('utf-8')), dest=response)

        # Verifica erros no processo de criação do PDF
        if pisa_status.err:
            return HttpResponse('Erro ao gerar PDF', status=500)

        return response

    except Exception as e:
        # Captura exceções e retorna erro
        return HttpResponse(f'Ocorreu um erro: {str(e)}', status=500)


def gerar_word_view(request, documento):


    # Cria um novo documento Word
    doc = Document()

    # Ajusta a formatação para Times New Roman e tamanho 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Processa o conteúdo HTML para extração do texto formatado
    soup = BeautifulSoup(documento.conteudo, "html.parser")

    # Itera sobre os elementos processados pelo BeautifulSoup
    for elemento in soup.children:
        if isinstance(elemento, str):
            # Adiciona texto puro (fora de tags)
            p = doc.add_paragraph(elemento.strip())
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
        else:
            # Tratamento para negrito e itálico
            if elemento.name == 'p':
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.5)  # Adiciona indentação na primeira linha
                for sub_elemento in elemento.children:
                    run = p.add_run(sub_elemento.text)

                    if sub_elemento.name == 'b':
                        run.bold = True
                    if sub_elemento.name == 'i':
                        run.italic = True
            elif elemento.name == 'h3':  # Trata títulos como negrito
                p = doc.add_paragraph(elemento.text.strip())
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)  # Um pouco maior para títulos
                p.paragraph_format.first_line_indent = Inches(0)
            elif elemento.name == 'h4':  # Trata subtítulos como negrito
                p = doc.add_paragraph(elemento.text.strip())
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)
                p.paragraph_format.first_line_indent = Inches(0)

    # Define o nome do arquivo e o tipo de conteúdo
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=documento_{documento.id}.docx'

    # Salva o arquivo Word no response
    doc.save(response)

    return response

def gerar_conteudo_contestacao(tipo_documento, dados_preenchimento):
    try:
        # Extração dos dados do caso para a contestação
        tipo_acao = dados_preenchimento.get('tipo_acao') or '[Tipo de Ação]'
        valor_causa = dados_preenchimento.get('valor_causa') or '[Valor Causa]'
        juizo_competente = dados_preenchimento.get('juizo_competente')
        fundamentacao_fatos = dados_preenchimento.get('fundamentacao_fatos')
        fundamentacao_direito = dados_preenchimento.get('fundamentacao_direito')
        dados_requerente = dados_preenchimento.get('dados_requerente') or '[Nome do Requerente]'
        dados_requerido = dados_preenchimento.get('dados_requerido') or '[Nome do Requerido]'
        processo_numero = dados_preenchimento.get('processo_numero') or '[Número do Processo]'
        
        # Parte 1: Buscar jurisprudências relevantes
        query = gerar_frase_pesquisa_gpt(fundamentacao_direito)
        jurisprudencias = buscar_jurisprudencias_bing(query)
        
        print(f"Gerando frase de pesquisa para o caso: '{query}'..")
        print("testeeeeeee", tipo_acao, valor_causa, juizo_competente, fundamentacao_fatos, fundamentacao_direito, dados_requerente, dados_requerido, processo_numero)
        # Exibir as jurisprudências encontradas
        #print(f"Jurisprudências encontradas: {jurisprudencias}")

        # Construção da mensagem para o modelo, incluindo jurisprudências reais
        mensagem = (
            f"Você é um advogado altamente especializado em direito civil, com mais de 20 anos de experiência \n"
            f"Voce precisa redigir uma Contestação. para uma ação do tipo: {tipo_acao}.\n"
            f"O valor da causa é de R$ {valor_causa}, e o juízo competente é {juizo_competente}.\n"
            f"A fundamentação dos fatos são o seguinte: {fundamentacao_fatos}.\n"
            f"Os dados conhecidos do requerente são: {dados_requerente}.\n"
            f"Os dados conhecidos do requerido são: {dados_requerido}.\n"
            f"O número do processo é: {processo_numero}.\n"
            f"As provas apresentadas são: {fundamentacao_direito}.\n"
            f"- Quando não forem fornecidos os dados das partes envolvidas, provas, processo, valores, não invente dados, utilize colchetes aguardando inclusão. Ex: [nome], [endereço].\n"
            f"\n\nInstruções adicionais:\n"
            f"- Inclua as jurisprudências relevantes da pesquisa e cite-as corretamente no formato de contestação jurídica em HTML.\n"
            f"- Inclua fundamentação legal apropriada, mencionando artigos do Código Civil, Código de Processo Civil ou outras legislações pertinentes.\n"
            f"- Caso seja citada a jurisprudência, utilize formatação em HTML, destacando trechos de legislação ou jurisprudência em <i>itálico</i>.\n"
            f"- Não inclua o link da pesquisa de jurisprudência no documento final. Formate as jurisprudências como citação profissional jurídica.\n"
            f"- O documento deve ser claro e objetivo, mantendo foco na defesa contra os fatos apresentados, evitando argumentos vagos ou irrelevantes.\n"
            f"- Todo o documento deve ser em HTML, sem usar ```html ou qualquer outra forma de marcação de código. Somente HTML puro.\n"
            f"- A estrutura deve conter seções como: <h3>DOS FATOS</h3>, <h3>DO DIREITO</h3>, <h3>DO PEDIDO</h3>.\n"
            f"- Use títulos como <h2> no início e <h3> ou <h4> ao longo do texto. Nunca use <h1>.\n"
            f"- Finalize com espaço para o nome do advogado, OAB e data.\n"
            f"- O documento deve ter clareza, fluidez e fundamentação robusta, mantendo a estrutura em HTML.\n"
            f"- O documento completo deve conter pelo menos 2000 palavras.\n"
            f"- As jurisprudências reais encontradas na web são as seguintes (use apenas as pertinentes):\n"
        )

        # Parte 2: Incluir as jurisprudências reais
        for idx, jurisprudencia in enumerate(jurisprudencias["webPages"]["value"][:3]):
            mensagem += f"{idx + 1}. Título: {jurisprudencia['name']}\n"
            mensagem += f"   URL: {jurisprudencia['url']}\n"
            mensagem += f"   Trecho: {jurisprudencia['snippet']}\n\n"

        mensagem += (
            "\nUse essas informações para compor a melhor contestação possível. "
            "Formate a citação em HTML e siga as instruções."
        )

        # Chamamos a API OpenAI usando streaming
        response = openai.chat.completions.create(
            model="chatgpt-4o-latest",  # Verifique se o modelo está correto
            messages=[
                {"role": "system", "content": "Você é um advogado especializado em direito civil."},
                {"role": "user", "content": mensagem}
            ],
            temperature=0.5,
            max_tokens=3640,
            stream=True
        )
        
        conteudo_completo = ""
        # Processar a resposta por partes
        for parte in response:
            if not parte.choices[0].delta == {}:
                print(str(parte.choices[0].delta.content), end="")
                conteudo_completo += str(parte.choices[0].delta.content)

        return conteudo_completo

    except Exception as e:
        print("Erro:", e)
        return "Erro: " + str(e)




